"""Document fixtures built in-process, with no PDF-writing dependency.

Word and Excel fixtures use ``python-docx`` and ``openpyxl``, which are already runtime
dependencies for *reading* those formats, so writing one costs nothing. PDF has no such
symmetry — ``pdfplumber`` only reads — and adding ``reportlab`` to the dev requirements to
produce a two-page test file would be a whole dependency for one fixture. So the PDF is
assembled here from raw objects.

It is about a hundred lines and it is worth them: the alternative is committing binary
fixtures, and a binary fixture is a file nobody can review, nobody can vary, and nobody can
tell has drifted from what the test claims it contains. :func:`ruled_table_pdf` in
particular produces genuine vector ruling lines, which is what pdfplumber's default table
strategy detects — so the table path is tested against a real table rather than mocked.
"""

from __future__ import annotations

import io

__all__ = [
    "text_pdf",
    "ruled_table_pdf",
    "scanned_pdf",
    "word_document",
    "workbook",
]


# --------------------------------------------------------------------------------- pdf


def _pdf(pages: list[tuple[list[tuple[float, float, str]], list[tuple]]]) -> bytes:
    """Assemble pages of positioned text and ruling lines into a PDF."""
    objects: dict[int, bytes] = {}
    count = len(pages)
    kids = " ".join(f"{4 + 2 * i} 0 R" for i in range(count))
    objects[1] = b"<< /Type /Catalog /Pages 2 0 R >>"
    objects[2] = f"<< /Type /Pages /Kids [{kids}] /Count {count} >>".encode()
    objects[3] = b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>"

    for index, (lines, rules) in enumerate(pages):
        stream = ""
        for x0, y0, x1, y1 in rules:
            stream += f"{x0} {y0} m {x1} {y1} l S\n"
        if lines:
            stream += "BT /F1 10 Tf\n"
            for x, y, text in lines:
                escaped = (
                    text.replace("\\", r"\\").replace("(", r"\(").replace(")", r"\)")
                )
                stream += f"1 0 0 1 {x} {y} Tm ({escaped}) Tj\n"
            stream += "ET\n"
        body = stream.encode()
        objects[4 + 2 * index] = (
            "<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            "/Resources << /Font << /F1 3 0 R >> >> "
            f"/Contents {5 + 2 * index} 0 R >>"
        ).encode()
        objects[5 + 2 * index] = (
            b"<< /Length "
            + str(len(body)).encode()
            + b" >>\nstream\n"
            + body
            + b"\nendstream"
        )

    out = bytearray(b"%PDF-1.4\n")
    offsets: dict[int, int] = {}
    for number in sorted(objects):
        offsets[number] = len(out)
        out += f"{number} 0 obj\n".encode() + objects[number] + b"\nendobj\n"

    start = len(out)
    top = max(objects) + 1
    out += f"xref\n0 {top}\n0000000000 65535 f \n".encode()
    for number in range(1, top):
        out += f"{offsets[number]:010d} 00000 n \n".encode()
    out += (
        f"trailer\n<< /Size {top} /Root 1 0 R >>\nstartxref\n{start}\n%%EOF\n"
    ).encode()
    return bytes(out)


def text_pdf(pages: list[list[str]]) -> bytes:
    """Prose only, one list of lines per page."""
    built = []
    for lines in pages:
        positioned = [(72.0, 720.0 - 16 * i, line) for i, line in enumerate(lines)]
        built.append((positioned, []))
    return _pdf(built)


def ruled_table_pdf(prose: list[str], rows: list[list[str]]) -> bytes:
    """Prose above a table with real ruling lines, on one page.

    Ruled rather than borderless because that is what pdfplumber's default strategy
    detects, and the default is what the extractor uses.
    """
    positioned = [(80.0, 700.0 - 16 * i, line) for i, line in enumerate(prose)]

    columns = [72.0 + 100.0 * i for i in range(len(rows[0]) + 1)]
    lines = [660.0 - 20.0 * i for i in range(len(rows) + 1)]
    rules = [(x, lines[-1], x, lines[0]) for x in columns]
    rules += [(columns[0], y, columns[-1], y) for y in lines]

    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            positioned.append((columns[c] + 5, lines[r] - 14, value))
    return _pdf([(positioned, rules)])


def scanned_pdf(pages: int = 1) -> bytes:
    """Pages with no text layer at all — what a scan looks like before OCR."""
    return _pdf([([], []) for _ in range(pages)])


# -------------------------------------------------------------------------------- word


def word_document(blocks: list[tuple[str, object]]) -> bytes:
    """Build a .docx from ``(kind, payload)`` pairs.

    ``kind`` is ``"heading1"``..``"heading3"``, ``"body"``, or ``"table"`` (payload is a
    list of rows). Order is preserved in the document body, which is the thing the
    extractor has to read correctly.
    """
    from docx import Document as WordDocument

    document = WordDocument()
    for kind, payload in blocks:
        if kind.startswith("heading"):
            document.add_heading(str(payload), level=int(kind[-1]))
        elif kind == "table":
            rows = payload  # type: ignore[assignment]
            table = document.add_table(rows=len(rows), cols=len(rows[0]))
            for r, row in enumerate(rows):
                for c, value in enumerate(row):
                    table.cell(r, c).text = str(value)
        else:
            document.add_paragraph(str(payload))

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


# ---------------------------------------------------------------------------- workbook


def workbook(sheets: dict[str, list[list[object]]], *, with_formula: bool = False) -> bytes:
    """Build an .xlsx. ``with_formula`` writes an uncached formula into the last cell.

    An uncached formula is how a file written by a tool that never calculated looks, and
    ``data_only=True`` reads it as empty — which is the trade the workbook extractor makes
    deliberately and the reason it warns rather than failing silently.
    """
    from openpyxl import Workbook

    book = Workbook()
    book.remove(book.active)
    for name, rows in sheets.items():
        sheet = book.create_sheet(title=name)
        for row in rows:
            sheet.append(row)
        if with_formula:
            sheet.cell(row=len(rows) + 1, column=1, value="=SUM(B2:B3)")

    buffer = io.BytesIO()
    book.save(buffer)
    return buffer.getvalue()
